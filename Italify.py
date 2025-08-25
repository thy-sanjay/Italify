import os
import re
from tkinter import Tk, filedialog
from docx import Document
from Bio import Entrez

# Set your email for NCBI Entrez
Entrez.email = "your.email@example.com"  # Replace with your actual email

def get_scientific_name(organism):
    try:
        handle = Entrez.esearch(db="taxonomy", term=organism, retmode="xml")
        records = Entrez.read(handle)
        handle.close()
        if records["Count"] == "0":
            return None
        tax_id = records["IdList"][0]
        summary_handle = Entrez.efetch(db="taxonomy", id=tax_id, retmode="xml")
        summary = Entrez.read(summary_handle)
        summary_handle.close()
        return summary[0]["ScientificName"]
    except Exception as e:
        print(f"Error fetching taxonomy info: {e}")
        return None

def italicize_scientific_names(text):
    organism_names = re.findall(r'\b[A-Z][a-z]* [a-z]+\b', text)
    for organism in set(organism_names):  # Use set to avoid re-processing
        scientific_name = get_scientific_name(organism)
        if scientific_name:
            text = re.sub(r'\b' + re.escape(organism) + r'\b', f'__ITALIC__{scientific_name}__END__', text)
    return text

def process_docx(input_path, output_path):
    doc = Document(input_path)
    for para in doc.paragraphs:
        if para.text.strip():
            new_text = italicize_scientific_names(para.text)
            para.clear()
            parts = re.split(r'(__ITALIC__.*?__END__)', new_text)
            for part in parts:
                if part.startswith('__ITALIC__') and part.endswith('__END__'):
                    run = para.add_run(part[10:-7])
                    run.italic = True
                else:
                    para.add_run(part)
    doc.save(output_path)

def convert_doc_to_docx(doc_path):
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_path)
        docx_path = doc_path + "x"
        doc.SaveAs(docx_path, FileFormat=16)
        doc.Close()
        word.Quit()
        return docx_path
    except Exception as e:
        print(f"Could not convert .doc to .docx: {e}")
        return None

def main():
    Tk().withdraw()
    file_path = filedialog.askopenfilename(
        title="Select a Word file",
        filetypes=[("Word files", "*.docx *.doc")]
    )

    if not file_path:
        print("No file selected.")
        return

    file_root, ext = os.path.splitext(file_path)

    if ext.lower() == ".doc":
        print("Converting .doc to .docx...")
        file_path = convert_doc_to_docx(file_path)
        if not file_path:
            print("Conversion failed.")
            return
        file_root, _ = os.path.splitext(file_path)

    default_name = os.path.basename(file_root) + "_itzd.docx"
    save_path = filedialog.asksaveasfilename(
        title="Save Italicized Document As",
        defaultextension=".docx",
        initialfile=default_name,
        filetypes=[("Word Document", "*.docx")]
    )

    if not save_path:
        print("Save cancelled.")
        return

    print("Processing document. This may take a while...")
    process_docx(file_path, save_path)
    print(f"Done! Italicized document saved as: {save_path}")

if __name__ == "__main__":
    main()
